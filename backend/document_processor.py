"""
文档解析模块

设计上明确两个真实约束:
1. 单文件最大 10MB —— 超大文件会拖慢同步解析响应,demo 现场卡顿比限制更致命。
   生产版本应改成异步处理 + 进度查询,这里先用同步保证 demo 稳定可控。
2. 扫描版 PDF(没有文字层,本质是图片)会被显式检测出来并报错,
   而不是悄悄返回空文本——那样会导致后续检索"假装成功但其实啥也没读到"。
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

import pypdf
import docx as docx_lib

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


class DocumentError(Exception):
    pass


@dataclass
class ParsedDocument:
    filename: str
    char_count: int
    text: str


def _decode_text_bytes(raw: bytes) -> str:
    """处理国内企业文档常见的编码问题:优先 utf-8,失败则尝试 gbk(老版 Office/Windows 导出常见)。"""
    for encoding in ("utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("文本编码无法识别(尝试了 utf-8/gbk/gb18030 均失败)")


def _parse_pdf(raw: bytes, filename: str) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(raw))
    except Exception as e:
        raise DocumentError(f"PDF 文件无法打开,可能已损坏: {e}")

    if reader.is_encrypted:
        raise DocumentError("PDF 文件已加密,当前版本不支持加密文档")

    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        # pypdf 解析某些中文字体(尤其是用了 CJK TrueType Collection 字体生成的PDF)时,
        # 会在每个字符之间插入空字符(\x00),不处理的话提取出来的内容会变成逐字乱码。
        page_text = page_text.replace("\x00", "")
        pages_text.append(page_text)
    full_text = "\n".join(pages_text).strip()

    if len(full_text) < 20:
        raise DocumentError(
            "未能从该 PDF 中提取到文字内容,这通常说明它是扫描版/图片版 PDF(没有文字层)。"
            "当前版本不支持 OCR,请上传文字版 PDF 或 txt/md/docx。"
        )
    return full_text


def _parse_docx(raw: bytes, filename: str) -> str:
    try:
        document = docx_lib.Document(io.BytesIO(raw))
    except Exception as e:
        raise DocumentError(f"docx 文件无法打开,可能已损坏: {e}")
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if len(text) < 20:
        raise DocumentError("该 docx 文档中几乎没有可提取的文字内容")
    return text


def parse_document(filename: str, raw: bytes) -> ParsedDocument:
    if len(raw) > MAX_FILE_SIZE_BYTES:
        raise DocumentError(f"文件大小 {len(raw) / 1024 / 1024:.1f}MB 超过 10MB 上限")
    if len(raw) == 0:
        raise DocumentError("文件内容为空")

    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentError(f"不支持的文件格式 {ext},目前支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if ext == ".pdf":
        text = _parse_pdf(raw, filename)
    elif ext == ".docx":
        text = _parse_docx(raw, filename)
    else:  # .txt / .md
        text = _decode_text_bytes(raw)

    text = text.strip()
    if len(text) < 20:
        raise DocumentError("解析后文档内容过短,可能不是有效文档")

    return ParsedDocument(filename=filename, char_count=len(text), text=text)


def chunk_text(text: str, chunk_size: int = 600, overlap: int = 50, min_chunk_size: int = 15) -> list[str]:
    """
    默认让每个语义段落(小标题+内容,空行分隔)独立成一个 chunk,这样检索精度最高——
    一个问题命中的就是真正相关的那一段,不会带出其他不相关板块的内容。
    只有两种情况例外:
      1. 段落小于 min_chunk_size(比如就一行短标题没内容),会并入下一段,避免出现无意义碎片。
      2. 段落本身超过 chunk_size(长篇大论的一段),会按句子边界二次切分并保留重叠。
    chunk_size 在这里更像是个"上限"而不是"目标值"。
    """
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    if not blocks:
        return []

    chunks: list[str] = []
    buf = ""
    for block in blocks:
        if len(block) > chunk_size:
            if buf:
                chunks.append(buf)
                buf = ""
            chunks.extend(_split_long_block(block, chunk_size, overlap))
        elif buf and len(buf) < min_chunk_size:
            buf = f"{buf}\n{block}"
        else:
            if buf:
                chunks.append(buf)
            buf = block
    if buf:
        chunks.append(buf)

    return chunks


def _split_long_block(block: str, chunk_size: int, overlap: int) -> list[str]:
    """对单个超长语义段落按句子边界二次切分,保留重叠避免句子被硬切断。"""
    sentences = [s for s in re.split(r"(?<=[。！？\n])", block) if s.strip()]
    pieces: list[str] = []
    buf = ""
    for s in sentences:
        if len(buf) + len(s) <= chunk_size:
            buf += s
        else:
            if buf:
                pieces.append(buf)
            tail = buf[-overlap:] if buf else ""
            buf = tail + s
    if buf:
        pieces.append(buf)
    return pieces
